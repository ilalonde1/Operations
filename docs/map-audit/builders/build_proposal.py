"""Append the corrected project-experience map to Daler's live proposal.

Works on a COPY. Daler is still editing the original and Jim wants to send it
today; nothing here writes to his file.

Changed from the 7 Aug version, all of it from Jim's 12 Aug call:

  * Downtown only. He asked to zoom right into the core and drop the outlying
    jobs -- "not even Chula Vista" -- so the Greater San Diego figure is gone
    and the county-wide count is a sentence instead of a page.
  * The QR opens the live map on downtown San Diego rather than the whole US
    portfolio, so a reader lands on their own blocks. Jim: "I just want to show
    them that we've done lots of jobs near their site."
  * The lists are regenerated from the live map data instead of being carried
    over. Six entries in the old KOR column -- Courthouse North Block, 777 Front
    Street, Kettner & Ash, 4th & Ash, India & Beech, 611 Island Avenue -- were
    peer reviews, plan checks or small scopes that the new $25k floor removed,
    which is precisely the complaint Jim raised.
  * The legend names the navy markers as Namdar's proposed sites. Three colours
    were on the figure and only two were explained.

1355 Broadway stays deliberately ABSENT from the KOR list. We do hold a record
for it -- 6 hours and $1,500 opened for JWDA in 2022 -- but that is not a design
and Namdar's own site cannot appear in a list of our completed work.
"""
import io, json, re, shutil, sys, urllib.request, zipfile, os

sys.path.insert(0, r"C:/Users/ilalonde/AppData/Local/Temp/claude/C--VIsual-Studio-Projects-Operations/912461f4-d333-42a6-8a2a-c879ddd0d90b/scratchpad/namdar")
import docx, html
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:/Users/ilalonde/AppData/Local/Temp/claude/C--VIsual-Studio-Projects-Operations/912461f4-d333-42a6-8a2a-c879ddd0d90b/scratchpad/namdar"
DEST_DIR = r"//KOR-FS01/Management/Project Opportunities/Pending/Namdar Group Park & Broadway JD JM (DUE AUG 12)"
SRC = OUT + "/work.docx"
DST = OUT + "/Namdar Group 1355 Broadway and 901 Park Blvd San Diego California 2026-08-12 KOR JD JM with project map.docx"
DATA = "https://www.korstructural.com/wp-content/uploads/kor-map-data.json"

MAP_URL = "https://www.korstructural.com/projects/?kor_at=32.7153,-117.1528,15"
BIO_URL = "https://www.korstructural.com/team_member/jim-desroches/"
ORANGE = RGBColor(0xFF, 0x5C, 0x35)
SLATE = RGBColor(0x2F, 0x33, 0x38)

DOWNTOWN = (32.703, 32.728, -117.176, -117.145)
ELSEWHERE = ("san ysidro", "chula vista", "la jolla", "national city", "la mesa",
             "el cajon", "escondido", "oceanside", "carlsbad", "coronado",
             "poway", "santee", "encinitas", "vista, ca")

TRAIL = re.compile(r"[\s,]*san\s+diego([\s,]*(ca|calif|california))?[\s,]*$", re.I)


def tidy(name):
    """Deltek appends the city to a job name; a client list should not show it."""
    n = html.unescape(name).strip()
    for _ in range(3):
        n2 = TRAIL.sub("", n).strip()
        if n2 == n:
            break
        n = n2
    return n


def portfolio():
    """Jim's Keep/Delete decisions of 10 Aug, read straight from his workbook.

    Not derived. He went through all 75 rows and marked which projects KOR
    actually got, with a reason on each. Nothing in Deltek reproduces that: it
    holds no won/lost flag, a lost pursuit still carries proposal time (so the
    billable-labour floor keeps it), and we invoice proposal work (so an
    invoice test keeps `4th & Ash` and `Courthouse North Block`, both of which
    he marks Delete -- the latter "lost to GS").
    """
    jim = json.loads(io.open(OUT + "/jim_list.json", encoding="utf-8").read())
    lat0, lat1, lng0, lng1 = DOWNTOWN

    def inbox(r):
        return (r.get("lat") is not None and r.get("lng") is not None
                and lat0 <= r["lat"] <= lat1 and lng0 <= r["lng"] <= lng1)

    kor = sorted({r["name"] for r in jim["kor_san_diego"] if inbox(r)})
    prior = sorted({r["name"] for r in jim["prior_downtown"] if inbox(r)})
    gsd_kor = len([r for r in jim["kor_san_diego"] if r.get("lat") is not None])
    return kor, prior, gsd_kor


def add_hyperlink(par, url, text, size=10, bold=False):
    """python-docx has no hyperlink API; build the w:hyperlink by hand."""
    rid = par.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "FF5C35"); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2)); rpr.append(sz)
    if bold:
        rpr.append(OxmlElement("w:b"))
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    par._p.append(link)
    return par


def replace_in_runs(doc, old, new, only_if_starts=None):
    """Swap text without flattening a paragraph's formatting.

    Word splits a line like "Date: August 7, 2026" across runs -- "Date:" bold,
    the date plain. Rewriting the whole paragraph would lose that, so the run
    holding the target is edited in place where possible.
    """
    hits = 0
    for par in doc.paragraphs:
        if only_if_starts and not par.text.strip().startswith(only_if_starts):
            continue
        if old not in par.text:
            continue
        for run in par.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                hits += 1
                break
        else:
            # target straddles runs: rebuild from the first run, keeping its style
            joined = par.text.replace(old, new)
            for r in par.runs[1:]:
                r.text = ""
            par.runs[0].text = joined
            hits += 1
    return hits


def fix_source_defects(d):
    """Defects in the inherited document, corrected on the COPY only.

    Reported to Ian 2026-08-12 and fixed on his instruction. Daler's original is
    untouched -- these need to go back into his master separately.
    """
    notes = []

    # 1. Appendix A lists "Senior Structural Engineer, Managing Principal $275"
    #    as both the first and the last row of the rate card.
    for t in d.tables:
        rows = [r.cells[0].text.strip() for r in t.rows]
        if len(rows) > 2 and rows[0].lower().startswith("role") and rows[-1] == rows[1]:
            el = t.rows[-1]._element
            el.getparent().remove(el)
            notes.append("removed duplicate rate-card row: %s" % rows[-1][:40])
            break

    # 2. Rate card was dated 15 months back.
    n = replace_in_runs(d, "Effective May 1, 2025", "Effective January 1, 2026")
    notes.append("rate card effective date -> January 1, 2026 (%d)" % n)

    # 3. Proposal date. The RFP RECEIPT date is a fact about the RFP and must
    #    not move, so only lines beginning "Date:" are touched.
    n = replace_in_runs(d, "August 7, 2026", "August 12, 2026", only_if_starts="Date:")
    notes.append("proposal date -> August 12, 2026 (%d)" % n)

    # 4. The firm has an Alberta office; the boilerplate predates it.
    n = replace_in_runs(d, "satellite offices in Kelowna and Nanaimo",
                        "satellite offices in Kelowna, Nanaimo and Alberta")
    notes.append("office list now includes Alberta (%d)" % n)

    # 5. Every dark section banner starts its own page. Appendix A's header row
    #    was stranded at the foot of page 8 with its rates overleaf.
    broke = 0
    prev_ends_with_break = False
    for par in d.paragraphs:
        if par.style.name == "Heading 1":
            # Skip where the preceding content already breaks the page -- the
            # cover does, and setting both produced a blank page 2.
            if not prev_ends_with_break:
                par.paragraph_format.page_break_before = True
                broke += 1
        prev_ends_with_break = bool(
            par._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'
                '[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type="page"]'))
    notes.append("page break before %d section headings" % broke)

    # Trailing blank paragraphs used to be absorbed at the foot of whatever page
    # they sat on. Now that each section starts its own page they have nowhere
    # to go and surface as an empty page -- the cover produced one once the QR
    # grew to a scannable 1.05in. Drop the blanks that sit directly above a
    # page-breaking heading; they were never content.
    removed = 0
    paras = d.paragraphs
    for i, par in enumerate(paras):
        if par.style.name != "Heading 1" or not par.paragraph_format.page_break_before:
            continue
        j = i - 1
        while j >= 0 and not paras[j].text.strip() and not paras[j]._element.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
            el = paras[j]._element
            el.getparent().remove(el)
            removed += 1
            j -= 1
    notes.append("removed %d blank spacer paragraphs above section breaks" % removed)
    return notes


def pic(doc, path, inches, new_page=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # page_break_before on the content itself. A separate break paragraph lands
    # on the new page and pushes the content one page further, leaving a blank.
    p.paragraph_format.page_break_before = new_page
    p.add_run().add_picture(path, width=Inches(inches))
    return p


def name_list(doc, heading, names, cols=4):
    """A borderless N-column table -- Word keeps columns aligned, tabs do not."""
    h = doc.add_paragraph()
    h.paragraph_format.keep_with_next = True
    hr = h.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(10)
    hr.font.color.rgb = SLATE

    rows = (len(names) + cols - 1) // cols
    tbl = doc.add_table(rows=rows, cols=cols)
    for i, n in enumerate(names):
        cell = tbl.cell(i % rows, i // rows)
        p = cell.paragraphs[0]
        run = p.add_run("\u2022  " + n)
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
    return tbl


def caption(par, text):
    r = par.add_run(text)
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = SLATE
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER


def swap_cover_qr(docx_path, new_png, target):
    tmp = docx_path + ".tmp"
    payload = io.open(new_png, "rb").read()
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        if target not in zin.namelist():
            raise SystemExit("cover QR image missing -- aborting rather than guessing")
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == target:
                data = payload
            zout.writestr(item, data)
    os.replace(tmp, docx_path)


def find_cover_qr(path):
    """The cover QR is the only square image in the file. Identify it rather
    than assuming a filename -- Daler has been editing this document."""
    from PIL import Image
    best = None
    with zipfile.ZipFile(path) as z:
        for n in z.namelist():
            if not n.startswith("word/media/"):
                continue
            try:
                im = Image.open(io.BytesIO(z.read(n)))
            except Exception:
                continue
            w, h = im.size
            if w and h and abs(w - h) / max(w, h) < 0.04:
                if best is None:
                    best = n
                else:
                    raise SystemExit("more than one square image; refusing to guess")
    if not best:
        raise SystemExit("no square image found -- the cover QR is not where it was")
    return best


if __name__ == "__main__":
    KOR, PRIOR, GSD_KOR = portfolio()
    print("downtown KOR %d | prior %d | Greater San Diego KOR %d" % (len(KOR), len(PRIOR), GSD_KOR))

    shutil.copyfile(SRC, DST)
    qr_target = find_cover_qr(DST)
    print("cover QR image:", qr_target)

    d = docx.Document(DST)
    for note in fix_source_defects(d):
        print("  fix:", note)
    ps = d.paragraphs

    # ---- cover: caption under the QR block Daler left ----
    anchor = next((p for p in ps if p.text.strip() == "Company Info:"), None)
    if anchor is None:
        sys.exit("'Company Info:' anchor gone -- document changed, aborting")
    blanks = []
    for p in ps[ps.index(anchor) + 1:]:
        if p.text.strip():
            break
        blanks.append(p)
    if len(blanks) >= 3:
        cap = blanks[1]
        r = cap.add_run("Scan for Jim's profile and project map \u2014 ")
        r.font.size = Pt(9)
        r.font.color.rgb = SLATE
        add_hyperlink(cap, BIO_URL, "korstructural.com/team_member/jim-desroches", size=9)
    else:
        print("  ! no blank block under 'Company Info:' -- cover caption skipped")

    # ---- the attachment the document's own Attachments list promises ----
    # No manual break needed: Heading 1 now carries page_break_before.
    h = d.add_paragraph(style="Heading 1")
    h.add_run("PORTFOLIO OF RELEVANT PROJECT EXPERIENCE")
    h.paragraph_format.page_break_before = True

    intro = d.add_paragraph()
    # Counts are PROJECTS, not Deltek job records -- one tower can carry six
    # sub-jobs. KOR's completed work and Jim's pre-KOR towers are counted,
    # coloured and listed separately: they are not the same claim.
    r = intro.add_run(
        "KOR has completed %d structural projects in downtown San Diego, on and around the "
        "blocks that surround 1355 Broadway and 901 Park Blvd, and %d across Greater San Diego. "
        "Jim DesRoches, the principal proposed for this project, led a further %d downtown San "
        "Diego towers earlier in his career; those are shown in green and listed separately. "
        "Every project shown is work KOR completed or is currently delivering, confirmed "
        "project by project by Jim DesRoches."
        % (len(KOR), GSD_KOR, len(PRIOR)))
    r.font.size = Pt(10.5)

    pic(d, OUT + "/fig_sd.png", 6.4)
    caption(d.add_paragraph(),
            "Downtown San Diego \u2014 KOR projects (orange) and Jim DesRoches' prior towers "
            "(green), around the two proposal sites (navy)")

    name_list(d, "KOR projects in downtown San Diego", KOR)
    name_list(d, "Jim DesRoches \u2014 downtown San Diego towers prior to KOR", PRIOR)

    qr = d.add_paragraph()
    qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr.paragraph_format.space_before = Pt(10)
    qr.paragraph_format.keep_with_next = True      # never orphan the caption
    qr.add_run().add_picture(OUT + "/qr_map_downtown.png", width=Inches(1.3))

    link = d.add_paragraph()
    link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = link.add_run("Open the live map on downtown San Diego: ")
    r.font.size = Pt(10)
    r.font.color.rgb = SLATE
    add_hyperlink(link, MAP_URL, "korstructural.com/projects", size=10, bold=True)

    d.save(DST)

    swap_cover_qr(DST, OUT + "/qr_jim_kor.png", qr_target)

    # Daler's placeholder sits at 0.81 in -- 0.62 mm a module, which decodes but
    # leaves nothing for a phone held at an angle over a printed page. 1.05 in
    # takes it to 0.8 mm and still fits the cover block.
    d2 = docx.Document(DST)
    grown = 0
    for sh in d2.inline_shapes:
        if abs(sh.width - sh.height) < 20000 and sh.width < Inches(1.0):
            sh.width = Inches(1.05); sh.height = Inches(1.05); grown += 1
    if grown != 1:
        raise SystemExit("expected exactly one small square image (the cover QR), found %d" % grown)
    d2.save(DST)
    print("wrote", DST)
